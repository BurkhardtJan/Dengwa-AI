import os
import shutil
from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
from uuid import UUID
from models import Media, MediaVocabulary, LanguageLearning, Chat
from schemas import VocabularyExtraction, MediaMetadataExtraction
from llm.prompts import build_vocab_extract_prompt, build_media_metadata_prompt, build_chat_title_prompt
from llm.client import call_llm
from llm.rag_service import embed_media
from services.vocabulary_service import get_or_create_vocab
from pypdf import PdfReader
from ebooklib import epub
import ebooklib
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from odf import text as odf_text
from odf import teletype
from odf.opendocument import load as odf_load
import logging

logger = logging.getLogger(__name__)
OCTET_STREAM_EXTENSIONS = {".txt", ".srt", ".vtt", ".md"}


def read_text_file(file_path: str) -> str:
    """Helper function to read text file"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read()


def extract_pdf_text(file_path: str) -> str | None:
    """Extracts pages from pdf file"""
    try:
        reader = PdfReader(file_path)
        text_parts = [
            text for page in reader.pages
            if (text := page.extract_text())
        ]
        return "\n\n".join(text_parts) if text_parts else None
    except Exception:
        return None


def extract_epub_text(file_path: str) -> str | None:
    """Extract text from epub file"""
    try:
        book = epub.read_epub(file_path)
        text_parts = []

        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text = soup.get_text(separator="\n", strip=True)
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts) if text_parts else None
    except Exception:
        return None


def extract_docx_text(file_path: str) -> str | None:
    """Extract text from docx"""
    try:
        doc = DocxDocument(file_path)
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # Auch Text aus Tabellen mitnehmen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n\n".join(text_parts) if text_parts else None
    except Exception:
        return None


def extract_odt_text(file_path: str) -> str | None:
    """Extract text from odt"""
    try:
        doc = odf_load(file_path)
        paragraphs = doc.getElementsByType(odf_text.P)
        text_parts = [
            "".join(node.data for node in p.childNodes if node.nodeType == node.TEXT_NODE)
            for p in paragraphs
        ]
        text_parts = [teletype.extractText(p) for p in paragraphs]

        return "\n\n".join(text_parts) if text_parts else None
    except Exception:
        return None


CONTENT_TYPE_EXTRACTORS: dict[str, callable] = {
    "application/x-subrip": read_text_file,
    "application/pdf": extract_pdf_text,
    "application/epub+zip": extract_epub_text,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx_text,
    "application/vnd.oasis.opendocument.text": extract_odt_text,
}


def extract_content(content_type: str, file_path: str) -> str | None:
    """Helper function to extract content from files"""
    if content_type.startswith("text/"):
        return read_text_file(file_path)

    if content_type in CONTENT_TYPE_EXTRACTORS:
        return CONTENT_TYPE_EXTRACTORS[content_type](file_path)

    if content_type == "application/octet-stream":
        ext = os.path.splitext(file_path)[1].lower()
        if ext in OCTET_STREAM_EXTENSIONS:
            return read_text_file(file_path)

    return None


def get_media_or_404(db: Session, media_id: UUID, user_id: UUID) -> Media:
    """Returns a Media record or raises 404."""
    media = (
        db.query(Media)
        .join(LanguageLearning, Media.learning_id == LanguageLearning.id)
        .filter(
            Media.id == media_id,
            LanguageLearning.user_id == user_id,
        )
        .first()
    )
    if not media:
        raise HTTPException(status_code=404, detail="Medium nicht gefunden")
    return media


def create_media_vocab(
        db: Session,
        media_id: UUID,
        learning_id: UUID,
        word: str,
        translation: str | None = None,
        context_sentence: str | None = None,
        language: str | None = None,
) -> MediaVocabulary:
    """
    Creates media vocabulary and if necessary normal vocabulary.
    """

    vocab = get_or_create_vocab(
        db=db,
        learning_id=learning_id,
        word=word,
        translation=translation,
        context_sentence=context_sentence,
        language=language
    )

    existing_link = (
        db.query(MediaVocabulary)
        .filter(
            MediaVocabulary.media_id == media_id,
            MediaVocabulary.vocabulary_id == vocab.id
        )
        .first()
    )

    if existing_link:
        return existing_link

    media_vocab_link = MediaVocabulary(
        media_id=media_id,
        vocabulary_id=vocab.id
    )

    db.add(media_vocab_link)
    db.commit()
    db.refresh(media_vocab_link)

    return media_vocab_link


def extract_and_save_vocabulary(db: Session, media: Media, provider, model) -> VocabularyExtraction:
    """
    Calls the LLM to extract vocabulary from a media item and persists results to DB.
    Returns the structured LLM response.
    """
    system_prompt = build_vocab_extract_prompt(media)
    messages = [{"role": "user", "content": "Gib zwischen 10 Vokabeln zurück"}]

    response_structured = call_llm(
        messages=messages,
        system_prompt=system_prompt,
        provider=provider,
        model=model,
        temperature=0.2,
        response_schema=VocabularyExtraction
    )

    for item in response_structured.vocabularies:
        create_media_vocab(
            db,
            media.id,
            media.learning_id,
            item.word,
            item.translation,
            item.context_sentence,
            media.language_learning.learning_language
        )

    return response_structured


def build_media_file_path(user_id: UUID, lan: str, media_id: UUID, ext: str) -> str:
    """Deterministic on-disk path for a medium's file, derived from its id.
    Never persisted — reconstructed on demand via get_media_disk_path()."""
    return os.path.join("uploads", str(user_id), lan, f"{media_id}{ext}")


def get_media_disk_path(media: Media) -> str:
    """Reconstructs a medium's on-disk file path from its id/extension
    plus its language_learning's user_id/learning_language."""
    learning = media.language_learning
    return build_media_file_path(learning.user_id, learning.learning_language, media.id, media.file_extension)


def save_uploaded_file(file: UploadFile, media_id: UUID, user_id: UUID, lan: str) -> str:
    """
    Saves an uploaded file to uploads/<user_id>/<lan>/<media_id><ext>.
    The filename is the Media's own id (generated by the caller before
    insert) plus the sanitized original extension — no separate UUID,
    and no user-supplied filename ever touches the path.
    Returns the file path.
    Raises HTTPException on write error.
    """
    ext = os.path.splitext(file.filename or "")[1].lower()
    ext = "".join(ch for ch in ext if ch.isalnum() or ch == ".")[:10]
    file_path = build_media_file_path(user_id, lan, media_id, ext)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Fehler beim Speichern der Datei: {str(e)}")

    return file_path


def create_media_record(db: Session, media_id: UUID, title: str, file: UploadFile, file_path: str,
                        learning_id: int) -> Media:
    """Creates and persists a Media DB record. Only the extension (not the
    full disk path) plus the original filename are stored for reference;
    the disk path itself is derived on demand via get_media_disk_path()."""
    media = Media(
        id=media_id,
        title=title,
        content_type=file.content_type,
        file_extension=os.path.splitext(file_path)[1],
        original_filename=file.filename,
        extracted_content=extract_content(file.content_type, file_path),
        learning_id=learning_id
    )
    db.add(media)
    db.commit()
    db.refresh(media)
    return media


def embed_media_safe(db: Session, media: Media) -> None:
    """
    Embeds a medium's content for RAG. Runs as a BackgroundTask — failures
    are logged instead of silently swallowed, so a broken embedding
    provider is visible without blocking or breaking the upload itself.
    """
    try:
        embed_media(db, media)
    except Exception:
        logger.exception("Failed to embed media %s", media.id)


def generate_media_metadata(db: Session, media_id: UUID, provider: str | None = None, model: str | None = None) -> None:
    """
    Summarizes a medium's content via LLM and persists summary/topics/
    genre/difficulty_estimate. Runs as a BackgroundTask after upload, so
    it doesn't delay the upload response.
    """
    media = db.get(Media, media_id)
    if not media or not media.extracted_content:
        return

    system_prompt = build_media_metadata_prompt(media)
    messages = [{"role": "user", "content": media.extracted_content}]

    result = call_llm(
        messages=messages,
        system_prompt=system_prompt,
        provider=provider,
        model=model,
        temperature=0.2,
        response_schema=MediaMetadataExtraction,
    )

    media.summary = result.summary
    media.topics = result.topics
    media.difficulty_estimate = result.difficulty_estimate
    media.genre = result.genre
    media.detected_language = result.detected_language
    db.commit()
